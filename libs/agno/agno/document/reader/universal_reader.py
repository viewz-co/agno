import asyncio
import importlib.util
from pathlib import Path
from typing import List, Union
from uuid import uuid4

from agno.document.base import Document
from agno.document.reader.base import Reader
from agno.utils.log import logger


class UniversalDocumentReader(Reader):
    """
    Universal document reader that replaces textract with modern, reliable alternatives.
    Supports PDF, DOCX, TXT, RTF, and other common document formats.
    """

    def __init__(self, chunk: bool = True):
        super().__init__(chunk=chunk)
        self._check_dependencies()

    def _check_dependencies(self):
        """Check if required dependencies are available"""
        self.available_readers = {}

        # Check for PDF readers
        if importlib.util.find_spec("pypdf") is not None:
            self.available_readers["pdf"] = "pypdf"
        elif importlib.util.find_spec("fitz") is not None:
            self.available_readers["pdf"] = "fitz"
        else:
            logger.warning("No PDF reader available. Install pypdf or PyMuPDF")

        # Check for DOCX support
        if importlib.util.find_spec("docx") is not None:
            self.available_readers["docx"] = "python-docx"
        else:
            logger.warning("python-docx not available. DOCX files won't be supported")

        # Check for OCR dependencies
        if (
            importlib.util.find_spec("pytesseract") is not None
            and importlib.util.find_spec("PIL") is not None
        ):
            self.available_readers["ocr"] = "pytesseract"
        else:
            logger.warning("OCR dependencies not available. Scanned documents won't be supported")

    def read(self, file_path: Union[str, Path]) -> List[Document]:
        """
        Read and extract text from various document formats

        Args:
            file_path: Path to the document file

        Returns:
            List of Document objects containing extracted text
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                logger.error(f"File not found: {file_path}")
                return []

            file_extension = file_path.suffix.lower()

            if file_extension == ".pdf":
                return self._read_pdf(file_path)
            elif file_extension in [".docx", ".doc"]:
                return self._read_docx(file_path)
            elif file_extension in [".txt", ".text"]:
                return self._read_text(file_path)
            elif file_extension == ".rtf":
                return self._read_rtf(file_path)
            elif file_extension in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
                return self._read_image_ocr(file_path)
            else:
                return self._read_fallback(file_path)

        except Exception as e:
            logger.error(f"Error reading document {file_path}: {e}")
            return []

    def _read_pdf(self, file_path: Path) -> List[Document]:
        """Extract text from PDF files"""
        if "pdf" not in self.available_readers:
            logger.error("No PDF reader available")
            return []

        documents = []

        try:
            if self.available_readers["pdf"] == "pypdf":
                from pypdf import PdfReader  # type: ignore

                reader = PdfReader(str(file_path))
                full_text = ""

                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    full_text += page_text + "\n"

                    if self.chunk:
                        doc = Document(
                            name=f"{file_path.stem}_page_{page_num + 1}",
                            id=str(uuid4()),
                            content=page_text,
                            meta_data={"page": page_num + 1, "source": str(file_path)},
                        )
                        documents.append(doc)

                if not self.chunk:
                    doc = Document(
                        name=file_path.stem,
                        id=str(uuid4()),
                        content=full_text,
                        meta_data={"source": str(file_path), "pages": len(reader.pages)},
                    )
                    documents.append(doc)

            elif self.available_readers["pdf"] == "fitz":
                import fitz  # type: ignore

                pdf_doc = fitz.open(str(file_path))
                full_text = ""

                for page_num in range(len(pdf_doc)):
                    page = pdf_doc.load_page(page_num)
                    page_text = page.get_text()
                    full_text += page_text + "\n"

                    if self.chunk:
                        document = Document(
                            name=f"{file_path.stem}_page_{page_num + 1}",
                            id=str(uuid4()),
                            content=page_text,
                            meta_data={"page": page_num + 1, "source": str(file_path)},
                        )
                        documents.append(document)

                if not self.chunk:
                    document = Document(
                        name=file_path.stem,
                        id=str(uuid4()),
                        content=full_text,
                        meta_data={"source": str(file_path), "pages": len(pdf_doc)},
                    )
                    documents.append(document)

                pdf_doc.close()

        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            return []

        return documents

    def _read_docx(self, file_path: Path) -> List[Document]:
        """Extract text from DOCX files"""
        if "docx" not in self.available_readers:
            logger.error("python-docx not available")
            return []

        try:
            import docx  # type: ignore

            doc = docx.Document(str(file_path))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = "\n".join(paragraphs)

            document = Document(
                name=file_path.stem,
                id=str(uuid4()),
                content=content,
                meta_data={"source": str(file_path), "format": "docx"},
            )

            documents = [document]

            if self.chunk:
                chunked_documents = []
                for document in documents:
                    chunked_documents.extend(self.chunk_document(document))
                return chunked_documents

            return documents

        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            return []

    def _read_text(self, file_path: Path) -> List[Document]:
        """Extract text from plain text files"""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            document = Document(
                name=file_path.stem,
                id=str(uuid4()),
                content=content,
                meta_data={"source": str(file_path), "format": "text"},
            )

            documents = [document]

            if self.chunk:
                chunked_documents = []
                for document in documents:
                    chunked_documents.extend(self.chunk_document(document))
                return chunked_documents

            return documents

        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return []

    def _read_rtf(self, file_path: Path) -> List[Document]:
        """Extract text from RTF files"""
        try:
            try:
                from striprtf.striprtf import rtf_to_text  # type: ignore

                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    rtf_content = f.read()

                content = rtf_to_text(rtf_content)

            except ImportError:
                logger.warning("striprtf not available, falling back to basic RTF parsing")
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()

                import re

                content = re.sub(r"\\[a-z]+\d*\s?", "", content)
                content = re.sub(r"[{}]", "", content)
                content = content.replace("\\", "")

            document = Document(
                name=file_path.stem,
                id=str(uuid4()),
                content=content,
                meta_data={"source": str(file_path), "format": "rtf"},
            )

            documents = [document]

            if self.chunk:
                chunked_documents = []
                for document in documents:
                    chunked_documents.extend(self.chunk_document(document))
                return chunked_documents

            return documents

        except Exception as e:
            logger.error(f"Error reading RTF file {file_path}: {e}")
            return []

    def _read_image_ocr(self, file_path: Path) -> List[Document]:
        """Extract text from images using OCR"""
        if "ocr" not in self.available_readers:
            logger.error("OCR dependencies not available")
            return []

        try:
            import pytesseract  # type: ignore
            from PIL import Image  # type: ignore

            image = Image.open(file_path)
            content = pytesseract.image_to_string(image)

            document = Document(
                name=file_path.stem,
                id=str(uuid4()),
                content=content,
                meta_data={"source": str(file_path), "format": "image_ocr"},
            )

            documents = [document]

            if self.chunk:
                chunked_documents = []
                for document in documents:
                    chunked_documents.extend(self.chunk_document(document))
                return chunked_documents

            return documents

        except Exception as e:
            logger.error(f"Error performing OCR on {file_path}: {e}")
            return []

    def _read_fallback(self, file_path: Path) -> List[Document]:
        """Fallback reader for unknown formats"""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            if not content.strip():
                with open(file_path, "rb") as f:
                    binary_content = f.read()
                    content = binary_content.decode("utf-8", errors="ignore")

            document = Document(
                name=file_path.stem,
                id=str(uuid4()),
                content=content,
                meta_data={"source": str(file_path), "format": "unknown"},
            )

            documents = [document]

            if self.chunk:
                chunked_documents = []
                for document in documents:
                    chunked_documents.extend(self.chunk_document(document))
                return chunked_documents

            return documents

        except Exception as e:
            logger.error(f"Error reading file {file_path} with fallback: {e}")
            return []

    async def async_read(self, file_path: Union[str, Path]) -> List[Document]:
        """Asynchronously read documents"""
        return await asyncio.to_thread(self.read, file_path)
